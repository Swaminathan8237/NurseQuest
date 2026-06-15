import os
import sys
import docx
from docx.shared import Pt, Inches

sys.stdout.reconfigure(encoding='utf-8')

def inspect_document(doc_path):
    print(f"\n==================================================")
    print(f"Inspecting Document: {os.path.basename(doc_path)}")
    print(f"==================================================")
    
    if not os.path.exists(doc_path):
        print(f"Error: File does not exist at {doc_path}")
        return
        
    doc = docx.Document(doc_path)
    
    # 1. Section margins
    print("\n--- Section Margins ---")
    for i, sec in enumerate(doc.sections):
        print(f"Section {i+1}:")
        print(f"  Top Margin:    {sec.top_margin.inches if sec.top_margin else 'Default'} in")
        print(f"  Bottom Margin: {sec.bottom_margin.inches if sec.bottom_margin else 'Default'} in")
        print(f"  Left Margin:   {sec.left_margin.inches if sec.left_margin else 'Default'} in")
        print(f"  Right Margin:  {sec.right_margin.inches if sec.right_margin else 'Default'} in")
        print(f"  Page Width:    {sec.page_width.inches if sec.page_width else 'Default'} in")
        print(f"  Page Height:   {sec.page_height.inches if sec.page_height else 'Default'} in")

    # 2. Check for unresolved placeholders or leftover notes
    print("\n--- Unresolved Placeholders / Leftover Note Markers ---")
    placeholders = ["[Note:", "[choose", "[specify", "TODO", "placeholder", "choose any one"]
    placeholder_matches = []
    
    # Check paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for p_marker in placeholders:
            if p_marker.lower() in text.lower():
                placeholder_matches.append(f"Para {i} (Style: {para.style.name}): '{text}'")
                
    # Check tables
    for t_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for p_marker in placeholders:
                    if p_marker.lower() in cell.text.lower():
                        placeholder_matches.append(f"Table {t_idx}, Row {row_idx}, Col {col_idx}: '{cell.text.strip()}'")
                        
    if placeholder_matches:
        for m in placeholder_matches:
            print(f"  [WARNING] Found potential placeholder: {m}")
    else:
        print("  No unresolved placeholders or note markers found.")

    # 3. Check for empty paragraphs at the end of pages or consecutive empty paragraphs
    print("\n--- Spacing & Empty Paragraphs Check ---")
    consecutive_empty = 0
    empty_count = 0
    total_paras = len(doc.paragraphs)
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            empty_count += 1
            consecutive_empty += 1
            if consecutive_empty >= 3:
                print(f"  [WARNING] Found {consecutive_empty} consecutive empty paragraphs around index {i}.")
        else:
            consecutive_empty = 0
            
    print(f"  Total Paragraphs: {total_paras}")
    print(f"  Empty Paragraphs: {empty_count} ({empty_count/total_paras*100:.1f}%)")

    # 4. Check formatting consistency (font family, sizes, alignment)
    print("\n--- Paragraph Formatting & Font Consistency Check ---")
    font_names = set()
    font_sizes = set()
    alignments = set()
    
    # We will sample formatting
    abnormal_formatting = []
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue
            
        align = para.alignment
        alignments.add(str(align))
        
        # Check fonts in runs
        for run in para.runs:
            if not run.text.strip():
                continue
            f_name = run.font.name
            f_size = run.font.size
            if f_name:
                font_names.add(f_name)
            if f_size:
                font_sizes.add(f_size.pt)
                
            # Check for font sizes outside standard ranges (typically 10-18 pt for body/headings)
            if f_size and (f_size.pt < 9 or f_size.pt > 36):
                abnormal_formatting.append(f"Para {i} run '{run.text[:20]}...' has unusual size: {f_size.pt} pt")
                
    print(f"  Font Names Used:  {list(font_names)}")
    print(f"  Font Sizes Used (pt): {sorted(list(font_sizes))}")
    print(f"  Alignments Used:   {list(alignments)}")
    if abnormal_formatting:
        print("  [WARNING] Abnormal formatting cases:")
        for case in abnormal_formatting[:5]:
            print(f"    - {case}")

    # 5. Table Layout & Cell Alignment Check
    print("\n--- Table Layout & Alignments ---")
    print(f"  Total Tables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        print(f"  Table {i+1}: {len(table.rows)} rows x {len(table.columns)} columns")
        # Check if table is aligned
        print(f"    Table Alignment: {table.alignment}")
        # Check cell alignment & widths
        width_mismatches = 0
        empty_cells = 0
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if not cell.text.strip():
                    empty_cells += 1
        print(f"    Empty Cells: {empty_cells}")

if __name__ == "__main__":
    # Let's inspect the files in the workspace
    project_root = r"d:\E0224033\Projeccts\Nursing_Education_Platform_gpt\Nursing_Education_Platform_gpt"
    
    filled_proposal = os.path.join(project_root, "CYB23IN201 Project proposal Internal - Filled.docx")
    inspect_document(filled_proposal)
    
    web_preview = os.path.join(project_root, "docs", "NurseQuest_Web_Preview.docx")
    inspect_document(web_preview)
    
    progress_update = os.path.join(project_root, "docs", "NurseQuest_Progress_Update.docx")
    inspect_document(progress_update)
