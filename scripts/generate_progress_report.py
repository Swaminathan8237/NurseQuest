"""
NurseQuest Progress Update — Word Document Generator
Generates a professional .docx documenting features added after the Web Preview.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
import shutil

# Paths
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
DEST_IMAGE_DIR = os.path.join(PROJECT_ROOT, "docs", "screenshots")

# The brain folder where screenshots were captured in the current session
BRAIN_DIR = r"C:\Users\Swaminathan\.gemini\antigravity-ide\brain\64768114-9960-4840-a0c6-528eebfa7365"

# Images mapping (source file name in brain directory -> destination name in screenshots directory)
IMAGE_MAPPING = {
    "landing_page_1780988945354.png": "landing_page.png",
    "module_manager_1780988982304.png": "module_manager.png",
    "module_view_student_1780989254172.png": "module_view_student.png",
    "slider_editor_1780989553048.png": "slider_editor.png",
    "matching_editor_1780989607616.png": "matching_editor.png"
}

def copy_images():
    print("--- 1. Copying Progress Update Screenshots ---")
    os.makedirs(DEST_IMAGE_DIR, exist_ok=True)
    
    # Also check other potential folders just in case
    alternative_brain = r"C:\Users\Swaminathan\.gemini\antigravity\brain\baa88481-06dd-459d-8851-f671d5c05c27"
    
    copied = 0
    for src_name, dest_name in IMAGE_MAPPING.items():
        src_path = os.path.join(BRAIN_DIR, src_name)
        if not os.path.exists(src_path):
            # Try alternative folder if not found in BRAIN_DIR
            src_path = os.path.join(alternative_brain, src_name)
            
        dest_path = os.path.join(DEST_IMAGE_DIR, dest_name)
        
        if os.path.exists(src_path):
            shutil.copy2(src_path, dest_path)
            print(f"  Copied {src_name} -> {dest_name} successfully.")
            copied += 1
        else:
            print(f"  Warning: {src_name} not found in brain directory.")
            
    print(f"Total screenshots copied: {copied}/{len(IMAGE_MAPPING)}")
    return copied > 0


# Helper: styled paragraph
def add_styled_paragraph(doc, text, font_size=11, bold=False, color=None, alignment=None, space_after=6, space_before=0, font_name="Calibri"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

# Helper: bullet lists
def add_bullet_item(doc, bold_prefix, text, font_size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    
    if bold_prefix:
        run_prefix = p.add_run(bold_prefix)
        run_prefix.font.name = "Calibri"
        run_prefix.bold = True
        run_prefix.font.size = Pt(font_size)
        
    run_text = p.add_run(text)
    run_text.font.name = "Calibri"
    run_text.font.size = Pt(font_size)
    return p

# Helper: image
def add_centered_image(doc, filename, width=Inches(5.5)):
    path = os.path.join(DEST_IMAGE_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(path, width=width)
    else:
        add_styled_paragraph(doc, f"[Screenshot: {filename} placeholder]", 10, color=(200, 0, 0), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)

# Helper: caption
def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    run.italic = True
    run.font.color.rgb = RGBColor(120, 120, 120)

def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ============================================================
    # COVER PAGE
    # ============================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(100)
    p.paragraph_format.space_after = Pt(12)
    # Use ASCII representation or text-based marker to prevent print crashes
    run = p.add_run("[Hospital Logo]")
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor(138, 79, 255)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("NURSEQUEST")
    run.font.name = "Calibri"
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(138, 79, 255)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("PROGRESS UPDATE & NEW FEATURES")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(70, 130, 180)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(160)
    run = p.add_run("Overview of Curriculum Organization, Advanced Quiz Question Types, and UX Enhancements Added Post-Web-Preview")
    run.font.name = "Calibri"
    run.font.size = Pt(12.5)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("JUNE 9, 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SREC · TECHNICAL PLATFORM DEVELOPMENT")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(70, 130, 180)

    doc.add_page_break()

    # ============================================================
    # PAGE 2: EXECUTIVE SUMMARY & WHAT IS NEW
    # ============================================================
    add_styled_paragraph(doc, "EXECUTIVE SUMMARY", font_size=18, bold=True, space_after=12, color=(138, 79, 255))

    add_styled_paragraph(
        doc,
        "Since the publication of the initial 'NurseQuest Web Preview' document, the platform has undergone "
        "significant architectural expansion and user-experience polishing. This Progress Update documents the "
        "features, pages, database schema modifications, and question formats implemented since then to transition the project "
        "from a basic quiz system into a full-scale, structured clinical training environment.",
        font_size=11, space_after=12
    )

    add_styled_paragraph(doc, "What is NOT in the original Web Preview document:", font_size=13, bold=True, space_before=12, space_after=8, color=(70, 130, 180))

    add_bullet_item(doc, "Interactive Landing Page & Guest Join Flow: ", "A pre-authorization homepage with theme toggles, detailed Host vs. Player experience grids, and an interactive phone mockup demonstrating friction-free 6-digit room PIN joins.")
    add_bullet_item(doc, "Curriculum-Based Learning Modules: ", "A core hierarchy grouping quizzes into clinical subjects (e.g., Fundamentals of Patient Care, Pharmacology) rather than simple text categories. Includes a dedicated Module Manager for teachers and Module View with progress indicators for students.")
    add_bullet_item(doc, "Advanced Question Formats (Slider & Matching Pairs): ", "Moving beyond multiple choice to support clinical vital range inputs (via Sliders) and term relationships (via Matching Pairs) both in single-player and Socket.io multiplayer.")
    add_bullet_item(doc, "Light / Dark Theme Support: ", "Full stylesheet token integration supporting real-time visual mode switches.")
    add_bullet_item(doc, "Expanded Database Schema: ", "Creation of the `modules` relational table and extension of the `questions` table with fields for slider constraints and matching pairs JSON arrays.")

    doc.add_page_break()

    # ============================================================
    # PAGE 3: LANDING PAGE & GUEST JOIN FLOW
    # ============================================================
    add_styled_paragraph(doc, "1. INTERACTIVE LANDING PAGE & GUEST JOIN FLOW", font_size=16, bold=True, space_after=8, color=(138, 79, 255))

    add_styled_paragraph(
        doc,
        "A major UX gap in the original web preview was the lack of an introductory page. The platform now features "
        "a high-impact landing page at the root route (`/`) which introduces the application to guests, explains "
        "the features, and provides a comparative review of Host vs. Player experiences.",
        font_size=11, space_after=10
    )

    add_styled_paragraph(doc, "Core Enhancements on this Page:", font_size=12, bold=True, space_before=6, space_after=6)
    add_bullet_item(doc, "Frictionless Guest Joining: ", "Students do not need to register. They can click 'Enter Code' directly on the landing page, opening a modal to submit a 6-digit room PIN and choose a nickname.")
    add_bullet_item(doc, "Profanity-Filtered Nickname Generator: ", "A casino-style random nickname generator is built into the join flow, combining safe prefixes and clinical terms (e.g., 'MightyPulse38') and checking against a blocked list of words to ensure classroom safety.")
    add_bullet_item(doc, "Theme Toggle: ", "A navbar controller allows users to switch between Light Mode and a sleek dark layout instantly.")
    add_bullet_item(doc, "Parallax Design: ", "Interactive background blur blobs track mouse movements, giving the landing page a premium, dynamic feel.")

    add_centered_image(doc, "landing_page.png", width=Inches(5.0))
    add_caption(doc, "Figure 1: Interactive Landing Page (/) featuring phone join mockup, theme switch, and statistics counters.")

    doc.add_page_break()

    # ============================================================
    # PAGE 4: LEARNING MODULES — TEACHER VIEW
    # ============================================================
    add_styled_paragraph(doc, "2. LEARNING MODULES — TEACHER VIEW", font_size=16, bold=True, space_after=8, color=(138, 79, 255))

    add_styled_paragraph(
        doc,
        "To structure learning pathways, we implemented a relational Learning Modules system. Instead of quizzes "
        "being loosely categorized by text strings, they are now child elements of defined Course Modules. "
        "Teachers manage this via the new Module Manager (`/modules`).",
        font_size=11, space_after=10
    )

    add_styled_paragraph(doc, "Module Manager Features:", font_size=12, bold=True, space_before=6, space_after=6)
    add_bullet_item(doc, "Module Creation & Metadata: ", "Teachers can define a Module Title and Description, select an appropriate icon (from 12 Material Icons), and choose a unique accent color (from 10 pre-curated color tokens) to color-coordinate the module's styling.")
    add_bullet_item(doc, "Drafting and Publishing: ", "Toggle switches determine whether a module is a Draft (invisible to students) or Published (live on student dashboards).")
    add_bullet_item(doc, "Drill-Down Quiz Management: ", "Expanding a module's row displays all quizzes assigned to that module, their draft/published states, question counts, and allows direct quiz editing or previewing.")

    add_centered_image(doc, "module_manager.png", width=Inches(5.5))
    add_caption(doc, "Figure 2: Teacher Module Manager (/modules) displaying the interactive module editor modal and quiz drill-down.")

    doc.add_page_break()

    # ============================================================
    # PAGE 5: LEARNING MODULES — STUDENT VIEW
    # ============================================================
    add_styled_paragraph(doc, "3. LEARNING MODULES — STUDENT VIEW", font_size=16, bold=True, space_after=8, color=(138, 79, 255))

    add_styled_paragraph(
        doc,
        "For students, modules provide a structured, gamified learning map. When a student clicks on a module "
        "from their dashboard, they are navigated to the dedicated Module View (`/modules/:id`), styled with "
        "the module's specific accent color.",
        font_size=11, space_after=10
    )

    add_styled_paragraph(doc, "Student Module Cockpit Features:", font_size=12, bold=True, space_before=6, space_after=6)
    add_bullet_item(doc, "Animated Progress Tracking: ", "A linear progress bar dynamically calculates and displays the percentage of quizzes completed within the module (e.g., '2/3 Quizzes Completed · 67%').")
    add_bullet_item(doc, "Visual Completion States: ", "Quizzes that have been attempted show a checkmark status icon and display the student's highest score. Unattempted quizzes show a numeric index and a 'Start' CTA.")
    add_bullet_item(doc, "Color-Themed UI Glows: ", "The entire UI adaptively updates its buttons, progress fills, and background glows to match the accent color designated by the teacher for that specific module.")

    add_centered_image(doc, "module_view_student.png", width=Inches(5.3))
    add_caption(doc, "Figure 3: Student Module View (/modules/:id) showing the color-coordinated progress bar, completion ticks, and quiz list.")

    doc.add_page_break()

    # ============================================================
    # PAGE 6: ADVANCED CLINICAL QUESTION TYPES
    # ============================================================
    add_styled_paragraph(doc, "4. ADVANCED CLINICAL QUESTION TYPES", font_size=16, bold=True, space_after=8, color=(138, 79, 255))

    add_styled_paragraph(
        doc,
        "To better simulate clinical scenarios, the Quiz Builder and Player have been expanded with two "
        "advanced question types: Sliders and Matching Pairs. These allow nursing students to interact "
        "with continuous numerical vital ranges and relational data.",
        font_size=11, space_after=10
    )

    add_styled_paragraph(doc, "Enhancement A: Slider Questions (Continuous Numeric Input)", font_size=13, bold=True, space_before=8, space_after=6, color=(70, 130, 180))
    add_styled_paragraph(
        doc,
        "Used for questions requiring numeric values (e.g. vital signs, drug dosages). Teachers configure the "
        "minimum value, maximum value, step size, unit (e.g., 'mmHg', 'bpm', '°C'), and the correct numeric answer. "
        "Students drag a custom-styled slider to submit their answer, which is checked against the database.",
        font_size=10.5, space_after=8
    )
    add_centered_image(doc, "slider_editor.png", width=Inches(4.8))
    add_caption(doc, "Figure 4: Quiz Builder — Slider Question configuration interface (setting min, max, step, unit, and answer).")

    add_styled_paragraph(doc, "Enhancement B: Matching Pairs (Relational Clinical Mapping)", font_size=13, bold=True, space_before=12, space_after=6, color=(70, 130, 180))
    add_styled_paragraph(
        doc,
        "Used to map symptoms to clinical diagnoses, or drugs to side effects. Teachers add list pairs. "
        "Students click left-side blocks and right-side blocks to form pairs, rendered with visual connection states. "
        "The scoring engine verifies all pairs on submit. Fully supported in both self-paced and live multiplayer.",
        font_size=10.5, space_after=8
    )
    add_centered_image(doc, "matching_editor.png", width=Inches(4.8))
    add_caption(doc, "Figure 5: Quiz Builder — Matching Pairs Question creation section.")

    doc.add_page_break()

    # ============================================================
    # PAGE 7: DATABASE SCHEMA & MIGRATIONS
    # ============================================================
    add_styled_paragraph(doc, "5. DATABASE SCHEMA UPDATES", font_size=16, bold=True, space_after=8, color=(138, 79, 255))

    add_styled_paragraph(
        doc,
        "To support these features, the SQLite schema was expanded from the original 9 tables to 10 tables, "
        "and new columns were appended to existing tables. Below is the updated relational structure.",
        font_size=11, space_after=12
    )

    # New Schema Table
    schema_data = [
        ["Table Name", "Structure & Primary Key", "Added Columns / Purpose"],
        ["modules", "id (TEXT PK)", "title (TEXT), description (TEXT), icon (TEXT), color (TEXT), order_index (INT), created_by (TEXT FK), is_published (INT). Groups quizzes into clinical categories."],
        ["quizzes", "id (TEXT PK)", "Added module_id (TEXT FK to modules). Links quizzes dynamically to modules rather than text categories."],
        ["questions", "id (TEXT PK)", "Added type check: allows 'slider' and 'matching'.\nAdded: slider_min (REAL), slider_max (REAL), slider_step (REAL), slider_unit (TEXT), matching_pairs (TEXT - JSON)."]
    ]

    table = doc.add_table(rows=len(schema_data) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    headers = ["Table Name", "Structure & Foreign Keys", "Added Columns / Purpose"]
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="8A4FFF"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data Rows
    for i, row_data in enumerate(schema_data, 1):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(9.5)
            if j == 0:
                run.bold = True
                run.font.name = "Consolas"
            if i % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0EAFF"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    add_styled_paragraph(doc, "Migration and Seed Integrations:", font_size=12, bold=True, space_before=18, space_after=6)
    add_bullet_item(doc, "db/migrate_modules.js: ", "A script to auto-generate the `modules` table, alter `quizzes` to add the foreign key, and cluster old loose quizzes into default modules.")
    add_bullet_item(doc, "db/migrate_new_types.js: ", "Re-creates the `questions` table schema with the slider-min/max/step/unit/matching-pairs columns and a CHECK constraint allowing 'slider' and 'matching' types.")
    add_bullet_item(doc, "db/seed.js: ", "Seeds three modules (Fundamentals, Anatomy & Physiology, Pharmacology) and populates multi-format questions (MCQ, Image, Video, Audio, Jumbles) with achievements and mock leaderboard data.")

    OUTPUT = os.path.join(PROJECT_ROOT, "docs", "NurseQuest_Progress_Update.docx")
    
    try:
        doc.save(OUTPUT)
        print(f"[OK] Document successfully saved to: {OUTPUT}")
    except PermissionError:
        print(f"[ERROR] Permission Denied: Could not save to {OUTPUT}. Please close the file if it is open in Microsoft Word.")
    except Exception as e:
        print(f"[ERROR] Failed to save document: {str(e)}")

if __name__ == "__main__":
    copy_images()
    # Generate the document
    print("Generating Progress Update Docx...")
    build_document()
