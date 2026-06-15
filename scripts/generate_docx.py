"""
NurseQuest Web Preview — Word Document Generator
Generates a professional .docx matching the Dental Web Preview style.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Paths
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
BASE = os.path.join(PROJECT_ROOT, "docs", "screenshots")
OUTPUT = os.path.join(PROJECT_ROOT, "docs", "NurseQuest_Web_Preview.docx")

SCREENSHOTS = {
    "auth": os.path.join(BASE, "auth_page_1777442625284.png"),
    "student": os.path.join(BASE, "student_dashboard_1777442700322.png"),
    "leaderboard": os.path.join(BASE, "leaderboard_page_1777442710539.png"),
    "teacher": os.path.join(BASE, "teacher_dashboard_1777442820808.png"),
    "quiz_builder": os.path.join(BASE, "quiz_builder_1777442829692.png"),
    "live_game": os.path.join(BASE, "live_game_1777442838805.png"),
}

doc = Document()

# ============================================================
# Page margins
# ============================================================
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ============================================================
# Helper: add colored paragraph
# ============================================================
def add_styled_paragraph(text, font_size=12, bold=False, color=None, alignment=None, space_after=6, space_before=0, font_name="Calibri"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_centered_image(path, width=Inches(5.5)):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run()
        run.add_picture(path, width=width)
    else:
        add_styled_paragraph(f"[Image not found: {os.path.basename(path)}]", 10, color=(200, 0, 0), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)

def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

# ============================================================
# COVER PAGE
# ============================================================

# Title (with spacing at top)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
p.paragraph_format.space_after = Pt(12)
run = p.add_run("🏥")
run.font.name = "Calibri"
run.font.size = Pt(48)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
run = p.add_run("NURSEQUEST")
run.font.name = "Calibri"
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(138, 79, 255)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
run = p.add_run("WEB PREVIEW")
run.font.name = "Calibri"
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(70, 130, 180)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(180)  # Large space after subtitle
run = p.add_run("Gamified Interactive Learning for Nursing Excellence")
run.font.name = "Calibri"
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
run = p.add_run("APRIL 29, 2026")
run.font.name = "Calibri"
run.font.size = Pt(12)
run.italic = True
run.font.color.rgb = RGBColor(100, 100, 100)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SREC\nSRINIVASAN")
run.font.name = "Calibri"
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(70, 130, 180)

doc.add_page_break()

# ============================================================
# PAGE 2: OVERVIEW & TECH STACK
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
run = p.add_run("NURSEQUEST WEBSITE")
run.font.name = "Calibri"
run.font.size = Pt(18)
run.bold = True
run.underline = True

add_styled_paragraph(
    "This is a gamified nursing education platform designed to help nursing students "
    "learn clinical concepts through interactive quizzes, real-time multiplayer sessions, "
    "and a competitive leaderboard system. The platform supports two roles — Teachers and Students — "
    "with role-based dashboards, XP progression, achievements, and avatar customization.",
    font_size=11, space_after=12
)

add_styled_paragraph("Tech Stack:", font_size=12, bold=True, space_after=6)

tech_items = [
    ("Frontend", "React 19, Vite, Tailwind CSS"),
    ("Backend", "Node.js with Express 5"),
    ("Database", "SQLite (via better-sqlite3)"),
    ("Real-time", "Socket.io (WebSocket)"),
    ("IDE", "VSCode and Antigravity"),
]

for label, value in tech_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(value)
    run2.font.size = Pt(11)

add_styled_paragraph("Flow with Images:", font_size=14, bold=True, space_before=18, space_after=12)

# Auth Page Screenshot
add_centered_image(SCREENSHOTS["auth"], width=Inches(4.5))
add_caption("Auth Page — Sign In / Sign Up (/auth)")

doc.add_page_break()

# ============================================================
# PAGE 3: FLOW DIAGRAM
# ============================================================
add_styled_paragraph("Application Flow Diagram:", font_size=14, bold=True, space_after=12)

# Create a text-based flow diagram since we can't embed Mermaid in docx
flow_text = """
                        ┌─────────────────┐
                        │   User Opens     │
                        │    NurseQuest    │
                        └────────┬────────┘
                                 │
                        ┌────────▼────────┐
                        │   Auth Page      │
                        │  /auth           │
                        │  Sign In/Sign Up │
                        └────────┬────────┘
                                 │
                    ┌────────────┼────────────┐
                    │                         │
           ┌───────▼───────┐        ┌────────▼────────┐
           │    TEACHER     │        │    STUDENT       │
           │   Dashboard    │        │   Dashboard      │
           │   /teacher     │        │   /student       │
           └───────┬───────┘        └────────┬────────┘
                   │                         │
        ┌──────────┼──────────┐    ┌─────────┼──────────┐
        │          │          │    │         │          │
   ┌────▼───┐ ┌───▼────┐ ┌──▼──┐ ┌──▼──┐ ┌──▼───┐ ┌──▼──┐
   │ Create │ │  Host  │ │View │ │Quiz │ │Lead- │ │Mini │
   │  Quiz  │ │  Live  │ │Board│ │Play │ │board │ │Game │
   │/builder│ │ /live  │ │     │ │/quiz│ │/lead │ │/mini│
   └────────┘ └────────┘ └─────┘ └─────┘ └──────┘ └─────┘
"""

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
run = p.add_run(flow_text)
run.font.name = "Consolas"
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(50, 50, 50)

# Flow description
flow_items = [
    "Teacher signs in → Teacher Dashboard → Create Quiz / Host Live / View Leaderboard",
    "Student signs in → (First time: Avatar Setup) → Student Dashboard",
    "Student Dashboard → Take Quiz / Join Live Game / View Leaderboard / Launch Mini-Game",
    "Quiz Completed → XP Earned → Level Up → Leaderboard Updated",
    "Live Game → Teacher hosts → Students join with 6-char code → Real-time sync via Socket.io",
]

for item in flow_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.size = Pt(10)

doc.add_page_break()

# ============================================================
# PAGE 4: STUDENT DASHBOARD
# ============================================================
add_styled_paragraph("Student Dashboard", font_size=16, bold=True, space_after=6, color=(138, 79, 255))

add_styled_paragraph(
    "The Student Dashboard displays personalized stats including Quizzes Taken, Average Score, "
    "Best Streak, Total XP, and the current nursing rank. Students can browse available quizzes, "
    "launch the IV Stabilization mini-game, and view top student rankings.",
    font_size=10, space_after=12
)

add_centered_image(SCREENSHOTS["student"], width=Inches(5.5))
add_caption("Student Dashboard (/student)")

doc.add_page_break()

# ============================================================
# PAGE 5: TEACHER DASHBOARD
# ============================================================
add_styled_paragraph("Teacher Dashboard", font_size=16, bold=True, space_after=6, color=(138, 79, 255))

add_styled_paragraph(
    "The Teacher Dashboard provides an overview of Total Students, Quizzes Created, Total Attempts, "
    "and Average Score. Teachers can create new quizzes, edit existing ones, publish/unpublish, "
    "host live sessions, and view individual student performance metrics.",
    font_size=10, space_after=12
)

add_centered_image(SCREENSHOTS["teacher"], width=Inches(5.5))
add_caption("Teacher Dashboard (/teacher)")

doc.add_page_break()

# ============================================================
# PAGE 6: QUIZ BUILDER
# ============================================================
add_styled_paragraph("Quiz Builder", font_size=16, bold=True, space_after=6, color=(138, 79, 255))

add_styled_paragraph(
    "The Quiz Builder allows teachers to create quizzes with 6 question types: "
    "Multiple Choice (MCQ), Image-Based, Video-Based, Audio-Based, Jumbled Letters, "
    "and Sequence Order. Each question includes configurable time limits, explanations, "
    "and point values.",
    font_size=10, space_after=12
)

add_centered_image(SCREENSHOTS["quiz_builder"], width=Inches(5.5))
add_caption("Quiz Builder — Create Quiz (/quiz-builder)")

doc.add_page_break()

# ============================================================
# PAGE 7: LEADERBOARD
# ============================================================
add_styled_paragraph("Leaderboard — The Arena", font_size=16, bold=True, space_after=6, color=(138, 79, 255))

add_styled_paragraph(
    "The Arena displays global rankings with a podium for Top 3 students. It shows each student's "
    "rank, avatar, level, quizzes taken, average score, and total XP. The sidebar displays the "
    "logged-in user's status, their target to overtake, rare achievements, and personal statistics.",
    font_size=10, space_after=12
)

add_centered_image(SCREENSHOTS["leaderboard"], width=Inches(5.5))
add_caption("Leaderboard — The Arena (/leaderboard)")

doc.add_page_break()

# ============================================================
# PAGE 8: LIVE GAME
# ============================================================
add_styled_paragraph("Live Game — Real-Time Multiplayer", font_size=16, bold=True, space_after=6, color=(138, 79, 255))

add_styled_paragraph(
    "The Live Game feature enables Kahoot-style real-time multiplayer quiz sessions. "
    "Teachers select a quiz and create a session with a 6-character join code. Students "
    "join using the code. Questions are synchronized across all participants using Socket.io, "
    "with time-based scoring and live leaderboard updates after each question.",
    font_size=10, space_after=12
)

add_centered_image(SCREENSHOTS["live_game"], width=Inches(5.5))
add_caption("Live Game — Host View (/live)")

doc.add_page_break()

# ============================================================
# PAGE 9: DATABASE SCHEMA
# ============================================================
add_styled_paragraph("Database Schema (SQLite)", font_size=16, bold=True, space_after=6, color=(138, 79, 255))

add_styled_paragraph(
    "The application uses SQLite with 9 relational tables. All data is stored locally "
    "in a single .db file using the better-sqlite3 driver for synchronous, fast queries.",
    font_size=10, space_after=12
)

# Create a styled table for the schema
table_data = [
    ["Table Name", "Purpose", "Key Columns"],
    ["users", "Student & Teacher accounts", "id, email, name, role, xp, level, streak, avatar_config"],
    ["quizzes", "Quiz metadata", "id, title, category, difficulty, time_per_question, is_published"],
    ["questions", "Individual questions", "id, quiz_id, type, question_text, options, correct_answer"],
    ["quiz_attempts", "Student quiz submissions", "id, quiz_id, user_id, score, correct_count, streak_max"],
    ["question_answers", "Per-question answers", "id, attempt_id, question_id, is_correct, points_earned"],
    ["live_sessions", "Multiplayer sessions", "id, quiz_id, host_id, join_code, status"],
    ["live_participants", "Session players", "id, session_id, user_id, score, streak, rank"],
    ["achievements", "Badge definitions", "id, name, requirement_type, requirement_value"],
    ["user_achievements", "Earned badges", "id, user_id, achievement_id, earned_at"],
]

table = doc.add_table(rows=len(table_data), cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Style the header row
for j, cell_text in enumerate(table_data[0]):
    cell = table.rows[0].cells[j]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(cell_text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(255, 255, 255)
    # Purple background for header
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="8A4FFF"/>')
    cell._tc.get_or_add_tcPr().append(shading)

# Style data rows
for i, row_data in enumerate(table_data[1:], 1):
    for j, cell_text in enumerate(row_data):
        cell = table.rows[i].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.font.size = Pt(9)
        if j == 0:
            run.bold = True
            run.font.name = "Consolas"
        # Alternate row color
        if i % 2 == 0:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0EAFF"/>')
            cell._tc.get_or_add_tcPr().append(shading)

# Add table borders
tbl = table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '</w:tblBorders>'
)
tblPr.append(borders)

doc.add_page_break()

# ============================================================
# PAGE 10: SCORING ENGINE
# ============================================================
add_styled_paragraph("Scoring & XP Engine", font_size=16, bold=True, space_after=6, color=(138, 79, 255))

add_styled_paragraph("Self-Paced Quiz Scoring:", font_size=12, bold=True, space_after=6)

scoring_items = [
    "Base Score: 1000 points (for a correct answer)",
    "Time Bonus: round(timeRemaining / totalTime × 500) → 0 to 500 pts",
    "Streak Bonus: min(streak, 5) × 100 → 0 to 500 pts",
    "Total per question: 1000 to 2000 points max",
    "Incorrect answer: 0 points, streak resets to 0",
]

for item in scoring_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.size = Pt(10)

add_styled_paragraph("Live Game Scoring (Kahoot-Style):", font_size=12, bold=True, space_before=18, space_after=6)

live_items = [
    "points = round(maxPoints × (1 − 0.5 × elapsed / totalTime))",
    "Fastest correct answer gets ~1000 pts, slowest gets ~500 pts",
    "Incorrect answer always scores 0",
]

for item in live_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.name = "Calibri"
    run.font.size = Pt(10)

add_styled_paragraph("XP Calculation:", font_size=12, bold=True, space_before=18, space_after=6)

xp_items = [
    "Base XP = score × 0.1",
    "100% Accuracy → +500 XP bonus",
    "≥ 80% Accuracy → +200 XP bonus",
    "≥ 60% Accuracy → +100 XP bonus",
    "Speed Ratio ≥ 0.9 → +150 XP bonus",
]

for item in xp_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.name = "Calibri"
    run.font.size = Pt(10)

add_styled_paragraph("Level Progression:", font_size=12, bold=True, space_before=18, space_after=6)

levels = [
    ("Level 1", "Nurse Intern", "0 XP", "🩺"),
    ("Level 2", "Junior Nurse", "1,000 XP", "💉"),
    ("Level 3", "Nurse", "3,000 XP", "🏥"),
    ("Level 4", "Senior Nurse", "6,000 XP", "⭐"),
    ("Level 5", "Head Nurse", "10,000 XP", "🌟"),
    ("Level 6", "Nurse Specialist", "15,000 XP", "💎"),
    ("Level 7", "Chief Nurse", "25,000 XP", "👑"),
]

level_table = doc.add_table(rows=len(levels) + 1, cols=4)
level_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
for j, header in enumerate(["Level", "Title", "Min XP", "Icon"]):
    cell = level_table.rows[0].cells[j]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(header)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(255, 255, 255)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="8A4FFF"/>')
    cell._tc.get_or_add_tcPr().append(shading)

for i, (level, title, xp, icon) in enumerate(levels, 1):
    for j, val in enumerate([level, title, xp, icon]):
        cell = level_table.rows[i].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(10)
        if i % 2 == 0:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0EAFF"/>')
            cell._tc.get_or_add_tcPr().append(shading)

# Borders for level table
tbl2 = level_table._tbl
tblPr2 = tbl2.tblPr if tbl2.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
borders2 = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '</w:tblBorders>'
)
tblPr2.append(borders2)

doc.add_page_break()

# ============================================================
# PAGE 11: POINTS TO NOTE
# ============================================================
add_styled_paragraph("POINTS TO NOTE:", font_size=14, bold=True, space_after=12)

notes = [
    "This platform uses a gamified approach inspired by Kahoot for nursing education. "
    "Quizzes support 6 question types including MCQ, Image-Based, Audio, Video, Jumbled Letters, and Sequence Order.",
    
    "Real-time multiplayer sessions use Socket.io WebSockets for synchronization. "
    "Teacher hosts the game, students join using a 6-character code.",
    
    "DB is Relational (i.e.) all information is stored in table structure using SQLite. "
    "9 tables handle users, quizzes, attempts, answers, live sessions, and achievements.",
    
    "The scoring engine uses a Kahoot-style time-bonus algorithm. Faster answers earn more points. "
    "Streak bonuses reward consecutive correct answers up to 5x.",
    
    "XP and Level progression system has 7 nursing ranks from 'Nurse Intern' (0 XP) to 'Chief Nurse' (25,000 XP).",

    "For further enhancement, AI-powered question generation, "
    "video/audio media playback, and cloud deployment can be integrated.",
]

for note in notes:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(note)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)

# ============================================================
# SAVE
# ============================================================
try:
    doc.save(OUTPUT)
    print(f"[OK] Document saved to: {OUTPUT}")
except PermissionError:
    print(f"[ERROR] Permission Denied: Could not save to {OUTPUT}. Please close the file if it is open in Microsoft Word or another application, then try again.")
except Exception as e:
    print(f"[ERROR] Failed to save document: {str(e)}")
