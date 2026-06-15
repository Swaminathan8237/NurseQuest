"""
Fill the CYB23IN201 Project Proposal Internal.docx with NurseQuest project content.
"""
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import os

INPUT_FILE = "CYB23IN201 Project proposal Internal.docx"
OUTPUT_FILE = "CYB23IN201 Project proposal Internal - Filled.docx"

doc = docx.Document(INPUT_FILE)


def delete_paragraph(paragraph):
    """Deletes a paragraph from the document XML tree."""
    p_element = paragraph._element
    parent = p_element.getparent()
    if parent is not None:
        parent.remove(p_element)


def set_paragraph_text(para, text, bold=False, font_size=None):
    """Clear a paragraph and set new text while preserving style."""
    # Keep the original style
    style = para.style
    alignment = para.alignment
    # Get formatting from the first run if it exists
    original_font_name = None
    original_font_size = None
    original_bold = None
    if para.runs:
        r = para.runs[0]
        original_font_name = r.font.name
        original_font_size = r.font.size
        original_bold = r.font.bold

    # Clear all runs
    for run in para.runs:
        run.text = ""
    
    # If no runs exist, add one
    if not para.runs:
        para.add_run("")
    
    para.runs[0].text = text
    if bold:
        para.runs[0].font.bold = True
    if font_size:
        para.runs[0].font.size = Pt(font_size)
    elif original_font_size:
        para.runs[0].font.size = original_font_size
    if original_font_name:
        para.runs[0].font.name = original_font_name


def add_paragraph_after(doc, ref_para, text, style=None, bold=False, font_size=11):
    """Insert a new paragraph after the reference paragraph and copy layout/indentation."""
    new_para = docx.oxml.OxmlElement('w:p')
    ref_para._element.addnext(new_para)
    new_paragraph = docx.text.paragraph.Paragraph(new_para, ref_para.part)
    if style:
        new_paragraph.style = style
    
    # Copy paragraph format to maintain alignment/indentation
    new_format = new_paragraph.paragraph_format
    ref_format = ref_para.paragraph_format
    new_format.left_indent = ref_format.left_indent
    new_format.right_indent = ref_format.right_indent
    new_format.first_line_indent = ref_format.first_line_indent
    new_format.space_before = ref_format.space_before
    new_format.space_after = ref_format.space_after
    new_format.line_spacing = ref_format.line_spacing
    new_format.alignment = ref_format.alignment

    run = new_paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    if bold:
        run.font.bold = True
    return new_paragraph


# === FILL EACH SECTION ===

paragraphs = doc.paragraphs

# --- Date ---
for p in paragraphs:
    if p.text.strip().startswith("Date:"):
        set_paragraph_text(p, "Date: 10-05-2026")
        break

# --- Project Domain ---
for p in paragraphs:
    if "Project Domain" in p.text and "Note:" in p.text:
        set_paragraph_text(p, "Project Domain: Full Stack Development")
        break

# --- Remove the "Choose any one" line and the list items ---
paras_to_delete = []
for p in paragraphs:
    if p.text.strip() == "Choose any one from the following":
        paras_to_delete.append(p)
    if p.style.name == "List Paragraph" and p.text.strip() in [
        "Cyber Security", "Internet of Things (IoT)", "Cloud Computing",
        "Full stack Development", "Mobile Application Development"
    ]:
        paras_to_delete.append(p)

for p in paras_to_delete:
    delete_paragraph(p)

# --- Project Title ---
for p in paragraphs:
    if p.text.strip().startswith("Project Title:"):
        set_paragraph_text(p, "Project Title: A Gamified Web-Based Quiz Platform for Interactive Nursing Education", bold=True)
        break
# Clear the note line below it
for p in paragraphs:
    if "[Note: Specify the exact working title" in p.text:
        delete_paragraph(p)
        break

# --- Computing Environment ---
computing_env = (
    "Computing Environment:\n"
    "• Operating System: Windows 10/11 (64-bit)\n"
    "• Processor: Intel Core i5 (or equivalent and above)\n"
    "• RAM: 8 GB (minimum)\n"
    "• Storage: 500 MB (application + dependencies)\n"
    "• Runtime: Node.js v18+ with NPM v9+\n"
    "• Database: SQLite 3 (embedded, serverless relational database via better-sqlite3)\n"
    "• IDE: Visual Studio Code\n"
    "• Build Tool: Vite v8.0 (frontend bundler and dev server)\n"
    "• Version Control: Git & GitHub\n"
    "• Browser: Google Chrome / Microsoft Edge"
)
for p in paragraphs:
    if "Computing Environment" in p.text and "Note:" in p.text:
        set_paragraph_text(p, computing_env)
        break

# --- Technologies Used ---
tech_used = (
    "Technologies used:\n"
    "• Languages: JavaScript (ES6+), HTML5, CSS3, SQL\n"
    "• Frontend: React 19, React Router v7, Vite, GSAP (animations), "
    "Three.js (3D avatars), Chart.js + react-chartjs-2 (analytics), "
    "Lottie-web (micro-animations), Canvas Confetti, Howler.js (audio), "
    "Socket.IO Client (real-time communication)\n"
    "• Backend: Node.js, Express 5 (REST API), Socket.IO (WebSocket server), "
    "better-sqlite3 (database), JWT (authentication), bcryptjs (password hashing), "
    "Multer (file uploads), UUID, CORS middleware\n"
    "• Dev Tools: Nodemon, ESLint, Concurrently"
)
for p in paragraphs:
    if "Technologies used:" in p.text and "Note:" in p.text:
        set_paragraph_text(p, tech_used)
        break

# --- Abstract ---
abstract = (
    "Abstract: "
    "Traditional nursing education relies heavily on passive learning approaches such as "
    "lectures and textbook reading, which often result in low student engagement and poor "
    "knowledge retention. This project presents NurseQuest, a gamified web-based quiz platform "
    "designed to transform nursing education into an interactive and engaging experience. The "
    "platform employs a full-stack architecture with a React-based frontend and a Node.js/Express "
    "backend, utilizing SQLite as the database. It supports two user roles — Teachers and Students "
    "— with distinct dashboards and capabilities. Teachers can create multimedia quizzes with "
    "multiple question types including multiple-choice, image-based, audio-based, video-based, "
    "jumbled-letter, and jumbled-sequence questions. Students can attempt quizzes independently "
    "or join real-time multiplayer quiz sessions powered by Socket.IO, competing on a live "
    "leaderboard with Kahoot-style scoring that rewards both accuracy and speed. The platform "
    "incorporates gamification elements such as experience points (XP), levelling, streaks, "
    "achievements, and customizable 3D avatars to foster motivation and sustained participation. "
    "Performance analytics with visual charts help both students and teachers track progress. "
    "NurseQuest demonstrates that gamification and real-time interactivity can significantly "
    "enhance the learning experience in nursing education."
)
for p in paragraphs:
    if "Abstract:" in p.text and "Note:" in p.text:
        set_paragraph_text(p, abstract)
        break

# --- Objectives ---
objectives = {
    "Objective 1:": (
        "Objective 1: To design and develop a full-stack gamified web application using React, "
        "Node.js, Express, and SQLite with role-based access control for teachers and students, "
        "featuring interactive dashboards, modern UI with animations, and a comprehensive quiz "
        "management system for the nursing education domain."
    ),
    "Objective 2:": (
        "Objective 2: To implement diverse and multimedia-rich questionnaire types — including "
        "Multiple Choice (MCQ), image-based, video-based, audio-based, jumbled-letter, and "
        "jumbled-sequence questions — enabling teachers to create varied and engaging assessments "
        "that go beyond traditional text-only quizzes."
    ),
    "Objective 3:": (
        "Objective 3: To integrate nursing-specific learning features such as categorized clinical "
        "quizzes (e.g., Anatomy, Pharmacology, Patient Care), unit and module-based question "
        "organization, medical scenario-based assessments, and domain-relevant gamification "
        "elements (XP, levels, streaks, badges) to reinforce nursing knowledge retention."
    ),
    "Objective 4:": (
        "Objective 4: To enable real-time, multi-device accessible multiplayer quiz sessions using "
        "Socket.IO and WebSocket communication, allowing students to join live competitive quizzes "
        "from any device (desktop, tablet, or mobile) simultaneously with instant scoring, live "
        "leaderboards, and synchronized question delivery."
    ),
}

for p in paragraphs:
    for key, value in objectives.items():
        if p.text.strip() == key:
            set_paragraph_text(p, value)
            break

# Clear the Note in the Project Objective heading
for p in paragraphs:
    if "Project Objective:" in p.text and "[Note:" in p.text:
        set_paragraph_text(p, "Project Objective:")
        break

# Objective 4 label in template -> now Objective 5 (production/deployment)
# The template only has Obj 1-4, so we repurpose Obj 4 and add Obj 5
# Actually, let me check: the template has Obj 1, 2, 3, 4. We need 5.
# Let me find Obj 4 position and add Obj 5 after it.

obj5_text = (
    "Objective 5: To ensure the platform is production-ready by implementing robust security "
    "measures (JWT authentication, bcrypt password hashing, input validation), optimizing for "
    "scalability through a single-server deployment architecture, and deploying the application "
    "for real-world accessibility."
)

# Find the paragraph for "Objective 4:" (now filled) and insert Obj 5 after it
for i, p in enumerate(paragraphs):
    if p.text.strip().startswith("Objective 4:"):
        new_p = add_paragraph_after(doc, p, obj5_text, style=p.style)
        break

# --- Methodology ---
methodology = (
    "Methodology: "
    "The project follows an Agile development methodology with iterative design and development "
    "cycles. The system adopts a client-server architecture with a React Single-Page Application "
    "(SPA) frontend and a Node.js/Express REST API backend communicating via HTTP and WebSocket "
    "protocols. The workflow consists of six modules: (1) Authentication Module — Users register "
    "and log in via JWT-based token authentication with bcrypt-hashed passwords and role-based "
    "access control; (2) Quiz Management Module — Teachers create quizzes through an interactive "
    "Quiz Builder with multiple question types and media attachments, stored in SQLite; "
    "(3) Quiz Attempt Module — Students browse and attempt published quizzes independently with "
    "responses scored and recorded for analytics; (4) Live Multiplayer Module — Teachers host "
    "real-time sessions generating a unique join code, students connect via Socket.IO, questions "
    "are broadcasted with countdown timers, and scores are calculated using a time-weighted "
    "Kahoot-style algorithm; (5) Gamification Engine — XP, levels, streaks, and achievement "
    "badges are computed and awarded based on performance; (6) Analytics Dashboard — Chart.js "
    "visualizations present performance trends, accuracy rates, and comparative analytics. "
    "A single-server deployment architecture serves both the API and static frontend files."
)
for p in paragraphs:
    if "Methodology:" in p.text and "Note:" in p.text:
        set_paragraph_text(p, methodology)
        break

# --- Deliverables ---
deliverables = (
    "Project outcome or Deliverables:\n"
    "1. Interactive Web Application — A responsive single-page application (SPA) with distinct "
    "interfaces for teachers and students, featuring modern UI design with animations, dark/light "
    "theme support, and 3D avatar customization.\n"
    "2. Quiz Builder Tool — A comprehensive quiz creation tool allowing teachers to build quizzes "
    "with six different question types (MCQ, image, video, audio, jumbled-letter, jumbled-sequence), "
    "set difficulty levels, assign categories, and publish quizzes.\n"
    "3. Real-Time Multiplayer Quiz Engine — A live, competitive quiz system where students join "
    "sessions via a code, answer questions in real-time against a timer, and compete on a dynamic "
    "leaderboard with Kahoot-style scoring.\n"
    "4. Gamification System — An XP, levelling, streak, and achievement badge system that "
    "incentivizes consistent learning and rewards performance milestones.\n"
    "5. Analytics & Reporting Dashboards — Visual performance dashboards showing quiz history, "
    "accuracy trends, leaderboard standings, and question-level analysis.\n"
    "6. Deployed Production Application — A fully deployed, security-hardened application with "
    "complete source code, database schema, and documentation."
)
for p in paragraphs:
    if "Project outcome or Deliverables:" in p.text:
        set_paragraph_text(p, deliverables)
        break
# Clear the note below
for p in paragraphs:
    if "[Note: Write the deliverables" in p.text:
        delete_paragraph(p)
        break

# --- References ---
references = [
    '1. Deterding, S., Dixon, D., Khaled, R. and Nacke, L., 2011. "From game design elements to gamefulness: Defining gamification." Proceedings of the 15th International Academic MindTrek Conference, pp. 9-15. ACM. DOI: 10.1145/2181037.2181040',
    '2. Brigham, T.J., 2015. "An introduction to gamification: Adding game elements for engagement." Medical Reference Services Quarterly, 34(4), pp. 471-480. DOI: 10.1080/02763869.2015.1082385',
    '3. Boctor, L., 2013. "Active-learning strategies: The use of a game to reinforce learning in nursing education." Nurse Education in Practice, 13(2), pp. 96-100. DOI: 10.1016/j.nepr.2012.07.010',
    '4. Wang, A.I. and Tahir, R., 2020. "The effect of using Kahoot! for learning - A literature review." Computers & Education, 149, p. 103818. DOI: 10.1016/j.compedu.2020.103818',
    '5. Tilton, K.J., Tiffany, J. and Hoglund, B.A., 2015. "Non-simulation game-based activities as teaching strategies for nursing students." Clinical Simulation in Nursing, 11(7), pp. 328-333. DOI: 10.1016/j.ecns.2015.04.005',
]

ref_idx = 0
for p in paragraphs:
    text = p.text.strip()
    if ref_idx < len(references) and text in ["1.", "2.", "3.", "4.", "5."]:
        set_paragraph_text(p, references[ref_idx])
        ref_idx += 1

# --- Clear the References note ---
for p in paragraphs:
    if "[Note: List minimum of five reference" in p.text:
        # This is embedded in the "References:" paragraph itself
        set_paragraph_text(p, "References:")
        break

# Save
doc.save(OUTPUT_FILE)
print(f"\n[OK] Filled proposal saved as: {OUTPUT_FILE}")
print(f"     Location: {os.path.abspath(OUTPUT_FILE)}")
