"""
Generate Live Game Analytics technical report as a Word document.
Output: docs/Live_Game_Analytics_Technical_Report.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
OUTPUT = os.path.join(PROJECT_ROOT, "docs", "Live_Game_Analytics_Technical_Report.docx")


def add_styled_paragraph(doc, text, font_size=11, bold=False, color=None, alignment=None,
                         space_after=6, space_before=0, font_name="Calibri"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 13, 3: 11.5}
    colors = {1: (0, 102, 153), 2: (51, 51, 51), 3: (68, 68, 68)}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = RGBColor(*colors.get(level, (0, 0, 0)))
    p.paragraph_format.space_before = Pt(12 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "Calibri"
        r1.font.size = Pt(10.5)
        r2 = p.add_run(text)
        r2.font.name = "Calibri"
        r2.font.size = Pt(10.5)
    else:
        run = p.runs[0] if p.runs else p.add_run(text)
        if not p.runs:
            run.text = text
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table


def build_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Title page
    add_styled_paragraph(doc, "NurseQuest / SkillQuest", font_size=14, bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_paragraph(doc, "Live Game Metrics & Data Recording", font_size=20, bold=True,
                         color=(0, 102, 153), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_styled_paragraph(doc, "Technical Report — What, Why, and How", font_size=12,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_styled_paragraph(doc, "Scope: Per-student live multiplayer game analytics", font_size=11,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_paragraph(doc, "Status: Implemented and production-ready", font_size=11,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary", 1)
    add_styled_paragraph(doc,
        "We built a persistent analytics layer for live multiplayer games so administrators can "
        "review each registered student's live-game performance with the same depth as solo quiz attempts.")
    add_styled_paragraph(doc,
        "Before this work, live games ran entirely in memory via Socket.IO. Scores appeared on the "
        "leaderboard during the session, but nothing was stored for later admin review. Solo quizzes "
        "already had quiz_attempts and question_answers; live games had only live_sessions and live_participants.")
    add_styled_paragraph(doc,
        "The solution mirrors the solo-quiz model with three new database tables, persists data when "
        "a game ends, exposes three admin API endpoints, and adds a Live Analytics mode in the Admin Dashboard.")

    # 2. Problem Statement
    add_heading(doc, "2. Problem Statement (Why)", 1)
    add_heading(doc, "2.1 What Was Missing", 2)
    add_table(doc,
        ["Capability", "Solo Quiz", "Live Game (Before)"],
        [
            ["Per-student attempt record", "Yes — quiz_attempts", "No"],
            ["Per-question answer record", "Yes — question_answers", "No"],
            ["Selection / hesitation tracking", "Yes (5-state status)", "No"],
            ["Admin drill-down report", "Yes — unit report", "No"],
            ["Post-session analytics", "Yes", "Lost when session ended"],
        ])

    add_heading(doc, "2.2 Design Goals", 2)
    add_bullet(doc, " parity with solo quiz analytics — same 5-state answer status and similar metrics.")
    add_bullet(doc, " minimal disruption to real-time gameplay — DB writes happen once at game-over.")
    add_bullet(doc, " admin-only visibility — students see simplified correct/incorrect; admins see full analytics.")
    add_bullet(doc, " data integrity — atomic writes, idempotent persistence, safe cleanup on user deletion.")

    # 3. Architecture
    add_heading(doc, "3. Architecture Overview", 1)
    add_styled_paragraph(doc, "Data flow:", bold=True, font_size=11)
    add_styled_paragraph(doc,
        "During game: LiveGame.jsx → Socket.IO (socket.js) → in-memory liveSessions Map\n"
        "At game over: persistLiveGameResults() → PostgreSQL transaction → live_game_attempts, "
        "live_game_answers, live_answer_selections\n"
        "Admin review: AdminDashboard.jsx → admin.js API routes → analytics.js formulas",
        font_size=10.5, space_after=8)
    add_styled_paragraph(doc,
        "Key principle: Real-time scoring stays in Socket.IO memory. Persistence is deferred until "
        "game-over, when all participant data is written in a single database transaction.")

    # 4. Database Schema Changes
    add_heading(doc, "4. Database Schema Changes", 1)
    add_styled_paragraph(doc,
        "This feature adds new persistence structures to PostgreSQL (Supabase). "
        "No existing tables were dropped or renamed. All changes are defined in "
        "backend/db/schema.sql and applied automatically on backend startup via initializeDB().")
    add_styled_paragraph(doc,
        "Summary: 3 new tables, 4 new indexes, 2 new columns (via ALTER TABLE), "
        "1 foreign-key constraint update, 0 table deletions.",
        bold=True, space_after=8)

    add_heading(doc, "4.1 New Tables Added", 2)
    add_table(doc,
        ["Table", "Purpose", "Mirrors"],
        [
            ["live_game_attempts", "One row per student per completed live session", "quiz_attempts"],
            ["live_game_answers", "One row per question per live attempt", "question_answers"],
            ["live_answer_selections", "Option-click trail before final submit", "N/A (live-only)"],
        ])

    add_heading(doc, "4.2 New Indexes Added", 2)
    add_table(doc,
        ["Index", "Table", "Column(s)"],
        [
            ["idx_lga_user", "live_game_attempts", "user_id"],
            ["idx_lga_session", "live_game_attempts", "session_id"],
            ["idx_lgans_attempt", "live_game_answers", "attempt_id"],
            ["idx_las_answer", "live_answer_selections", "answer_id"],
        ])

    add_heading(doc, "4.3 Columns Added (ALTER TABLE)", 2)
    add_table(doc,
        ["Table", "Column", "Purpose"],
        [
            ["live_game_answers", "status", "5-state outcome: correct, incorrect, selected_correct, selected_incorrect, not_answered"],
            ["live_answer_selections", "is_correct", "Whether each staged click matched the answer key"],
        ])
    add_styled_paragraph(doc,
        "Uses ADD COLUMN IF NOT EXISTS so existing production databases upgrade safely without data loss.")

    add_heading(doc, "4.4 Foreign-Key Constraint Update", 2)
    add_styled_paragraph(doc,
        "live_game_attempts.user_id was updated to reference users(id) ON DELETE CASCADE. "
        "When a student is permanently deleted, their live game attempts (and cascaded answers "
        "and selections) are removed automatically. executePermanentDeletion() in admin.js also "
        "explicitly deletes live_game_attempts before removing the user as a defensive measure.")

    add_heading(doc, "4.5 What Was NOT Changed", 2)
    add_bullet(doc, " No existing tables were dropped (users, quizzes, quiz_attempts, live_sessions, etc.).")
    add_bullet(doc, " No columns were removed from existing tables.")
    add_bullet(doc, " Solo-quiz tables (quiz_attempts, question_answers) were not modified for this feature.")
    add_bullet(doc, " A separate manual migration script is not required — schema.sql runs on every server start.")

    # 5. Data Model
    add_heading(doc, "5. Data Model (What We Store)", 1)

    add_heading(doc, "5.1 live_game_attempts", 2)
    add_styled_paragraph(doc,
        "One row per registered student per completed live session (mirrors quiz_attempts).")
    add_table(doc,
        ["Column", "Purpose"],
        [
            ["session_id", "Links to live_sessions"],
            ["user_id", "Registered student only"],
            ["final_score", "Total Kahoot-style points earned"],
            ["total_points", "Maximum possible points"],
            ["correct_count", "Number of correct answers"],
            ["total_questions", "Questions in the session"],
            ["max_streak", "Longest correct-answer streak"],
            ["total_time_ms", "Sum of response times"],
            ["final_rank", "Leaderboard position at game end"],
            ["completed_at", "Timestamp"],
        ])
    add_styled_paragraph(doc, "Constraint: UNIQUE(session_id, user_id) — one attempt per student per session.")

    add_heading(doc, "5.2 live_game_answers", 2)
    add_styled_paragraph(doc, "One row per question per attempt (mirrors question_answers).")
    add_table(doc,
        ["Column", "Purpose"],
        [
            ["attempt_id", "Parent attempt"],
            ["question_id", "Question reference"],
            ["question_index", "Order in session (0-based)"],
            ["final_answer", "Submitted answer (text/JSON)"],
            ["is_correct", "Graded correctness"],
            ["points_earned", "Kahoot-style score for this question"],
            ["response_ms", "Time to answer (milliseconds)"],
            ["is_timeout", "Server-side timeout finalization"],
            ["is_late", "Submitted after timer expired"],
            ["status", "5-state outcome (see Section 6.3)"],
        ])

    add_heading(doc, "5.3 live_answer_selections", 2)
    add_styled_paragraph(doc,
        "Selection trail for select-then-confirm question types: mcq, image, video, audio.")
    add_table(doc,
        ["Column", "Purpose"],
        [
            ["answer_id", "Parent answer row"],
            ["selection_order", "1, 2, 3… (each option click)"],
            ["selected_value", "Option chosen"],
            ["selected_at", "Unix timestamp (ms)"],
            ["elapsed_ms", "Time since question started"],
            ["is_correct", "Whether that selection matched the key"],
        ])
    add_styled_paragraph(doc, "Cap: 20 selections per question (abuse prevention).")

    add_heading(doc, "5.4 Referential Integrity", 2)
    add_bullet(doc, " live_game_attempts.session_id → live_sessions ON DELETE CASCADE")
    add_bullet(doc, " live_game_attempts.user_id → users ON DELETE CASCADE")
    add_bullet(doc, " live_game_answers.attempt_id → live_game_attempts ON DELETE CASCADE")
    add_bullet(doc, " live_answer_selections.answer_id → live_game_answers ON DELETE CASCADE")
    add_styled_paragraph(doc,
        "User permanent deletion explicitly removes live_game_attempts before deleting the user, "
        "preventing foreign-key violations.")

    # 6. Runtime collection
    add_heading(doc, "6. Runtime Data Collection (How — During Game)", 1)
    add_heading(doc, "6.1 In-Memory Participant State", 2)
    add_styled_paragraph(doc,
        "When a student joins (join-session), the server stores: score, streak, maxStreak, answers[], "
        "selectionTrails{}, totalResponseMs, joinedAt.")

    add_heading(doc, "6.2 Answer Submission (submit-answer)", 2)
    add_bullet(doc, " Final answer value and type-specific grading (MCQ, captcha, matching, slider, etc.)")
    add_bullet(doc, " Kahoot-style score via calculateLiveScoreKahootStyle()")
    add_bullet(doc, " Response time in milliseconds")
    add_bullet(doc, " committed flag — true for normal submit, false for timer-expiry staged selection")

    add_heading(doc, "6.3 Five-State Answer Status", 2)
    add_styled_paragraph(doc, "Aligned with solo quiz semantics (question_answers.status):")
    add_table(doc,
        ["Status", "Meaning"],
        [
            ["correct", "Committed answer, graded correct"],
            ["incorrect", "Committed answer, graded wrong"],
            ["selected_correct", "Timer expired with staged selection that would be correct"],
            ["selected_incorrect", "Timer expired with staged selection that would be wrong"],
            ["not_answered", "No answer and no selection trail"],
        ])
    add_styled_paragraph(doc,
        "Computed by computeLiveAnswerStatus() in backend/socket.js using the committed flag, "
        "final answer, and selection trail fallback for host-skip scenarios.")

    add_heading(doc, "6.4 Selection Trail (selection-change)", 2)
    add_styled_paragraph(doc,
        "For MCQ, image, video, and audio: first click selects (selection-change emitted); "
        "second click confirms (submit-answer with committed: true); timer expiry submits staged "
        "option with committed: false.")

    add_heading(doc, "6.5 Who Gets Persisted", 2)
    add_styled_paragraph(doc,
        "Only registered students (role = student) with non-guest IDs. Guest participants (guest_*) "
        "are excluded because they have no persistent user record.")

    # 7. Persistence
    add_heading(doc, "7. Persistence Pipeline (How — At Game Over)", 1)
    add_styled_paragraph(doc, "When the last question finishes:")
    add_bullet(doc, " Emit game-over with final rankings")
    add_bullet(doc, " Update live_sessions.status = finished, set ended_at")
    add_bullet(doc, " Call persistLiveGameResults(session, rankings)")
    add_styled_paragraph(doc, "Transaction properties:", bold=True)
    add_bullet(doc, " Atomic — all-or-nothing per game")
    add_bullet(doc, " Idempotent — duplicate game-over does not create duplicate attempts")
    add_bullet(doc, " Non-blocking — persistence runs async; UI unaffected if DB write fails (error logged)")

    # 8. Admin API
    add_heading(doc, "8. Admin Analytics API", 1)
    add_table(doc,
        ["Endpoint", "Purpose"],
        [
            ["GET /admin/students/:id/live-games", "List all live sessions for a student"],
            ["GET /admin/live-games/:attemptId/detail", "Question-by-question drill-down with selection trail"],
            ["GET /admin/students/:id/live-report", "Aggregated metrics, knowledge score, badges"],
        ])

    add_heading(doc, "8.1 Live Report Metrics", 2)
    add_bullet(doc, " Per session: accuracy, first-attempt accuracy, speed score, efficiency, time utilization")
    add_bullet(doc, " Knowledge Score + classification (Excellent / Good / Moderate / Poor / Very Poor)")
    add_bullet(doc, " Retention when same quiz played 2+ times")
    add_bullet(doc, " Overall aggregates across all live games")

    add_heading(doc, "8.2 Badges (8 total)", 2)
    add_table(doc,
        ["Badge", "Criteria"],
        [
            ["Accuracy", "Overall accuracy > 90%"],
            ["Speed", "Speed score ≥ 90"],
            ["Consistency", "≥ 3 games with 80%+ accuracy"],
            ["Perfect Score", "100% accuracy in a game"],
            ["Mastery", "All first selections correct in a game"],
            ["Streak", "Max streak ≥ 5"],
            ["Top Ranker", "Achieved 1st place"],
            ["Endurance", "≥ 10 games played"],
        ])

    add_heading(doc, "8.3 Separation from Unit Report", 2)
    add_styled_paragraph(doc,
        "The existing /admin/students/:id/report (unit report) explicitly excludes live games via "
        "meta.liveGamesExcluded: true. Live analytics use dedicated endpoints and UI mode.")

    # 9. Frontend
    add_heading(doc, "9. Admin Dashboard UI", 1)
    add_styled_paragraph(doc,
        "AdminDashboard.jsx adds an analytics mode toggle: Units (solo quiz) vs Live (live games). "
        "Live mode loads live-report, live-games list, and attempt detail on click. Status badges use "
        "the same 5-state map as solo quizzes.")

    # 10. Analytics engine
    add_heading(doc, "10. Analytics Engine (utils/analytics.js)", 1)
    add_table(doc,
        ["Function", "Purpose"],
        [
            ["accuracy(correct, total)", "(correct / total) × 100"],
            ["speedScore(expected, actual)", "(expected / actual) × 100, capped at 100"],
            ["efficiency(accuracy, avgResponse)", "accuracy / avgResponseTime"],
            ["retention(later, initial)", "(later / initial) × 100"],
            ["knowledgeScore({...})", "50% accuracy + 20% first-attempt + 15% speed + 15% retention"],
            ["classify(score)", "Maps 0–100 to knowledge level band"],
            ["sumField(rows, field)", "Safe numeric summation for report rollups"],
        ])

    # 11. Reliability fixes
    add_heading(doc, "11. Post-Implementation Reliability Fixes", 1)
    add_table(doc,
        ["Issue", "Impact", "Fix"],
        [
            ["FK violation on user deletion", "Permanent delete failed if student had live records",
             "ON DELETE CASCADE + explicit cleanup in executePermanentDeletion"],
            ["Video/audio timeout dropped selection", "Staged option lost on timer expiry",
             "Extended timeout auto-submit to video/audio types"],
            ["Wrong status for staged selections", "Timed-out questions stored as not_answered",
             "computeLiveAnswerStatus() uses committed flag + trail fallback"],
        ])

    # 12. Exclusions
    add_heading(doc, "12. What Is NOT Recorded", 1)
    add_table(doc,
        ["Excluded", "Reason"],
        [
            ["Guest players (guest_*)", "No persistent user ID"],
            ["Mid-game disconnects", "Only completed games trigger persistence"],
            ["Host/teacher participants", "Filtered to role = student"],
            ["In-progress sessions", "Data only in memory until game-over"],
            ["Live games in unit report", "Intentionally separate analytics tracks"],
        ])

    # 13. Conclusion
    add_heading(doc, "13. Conclusion", 1)
    add_styled_paragraph(doc,
        "Live game metrics and per-student data recording are fully implemented. The system records "
        "every registered student's live game attempt, per-question answers, and option-click trails; "
        "computes the same analytics metrics as solo quizzes; exposes three admin API endpoints and "
        "a dedicated Live Analytics UI mode; preserves real-time performance by deferring DB writes "
        "to game-over; and maintains data integrity through transactions, idempotency, cascade deletes, "
        "and 5-state status semantics.")

    add_styled_paragraph(doc, "— End of Report —", font_size=10,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24)

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(f"Generated: {path}")
